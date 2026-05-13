"""
report_agent.py — Gen report .docx ด้วย python-docx
Auto-triggered หลัง save_chatlog สำเร็จ
"""

import os
import json
import asyncio
import tempfile
from datetime import datetime, timezone, timedelta

import httpx
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload

TZ_THAI = timezone(timedelta(hours=7))

ANTHROPIC_API_KEY = os.environ["ANTHROPIC_API_KEY"]
ANTHROPIC_API_URL = "https://api.anthropic.com/v1/messages"
CLAUDE_MODEL = "claude-sonnet-4-20250514"
DRIVE_FOLDER_ID = os.environ["GOOGLE_DRIVE_FOLDER_ID"]
LINE_CHANNEL_ACCESS_TOKEN = os.environ["LINE_CHANNEL_ACCESS_TOKEN"]
PAYAT_USER_ID = "U86c03cd5153459d2dc9ce52adc608147"

SCOPES = [
    "https://www.googleapis.com/auth/drive",
    "https://www.googleapis.com/auth/drive.file",
    "https://www.googleapis.com/auth/spreadsheets",
]


# ─── Google Drive (OAuth — เหมือนน้องแพลน) ───────────────────────────────────
def _drive():
    from google.oauth2.credentials import Credentials
    from google.auth.transport.requests import Request

    # ไม่ใส่ scopes — ใช้ pattern เดียวกับน้องแพลน
    creds = Credentials(
        token=None,
        refresh_token=os.environ["GOOGLE_OAUTH_REFRESH_TOKEN"],
        client_id=os.environ["GOOGLE_OAUTH_CLIENT_ID"],
        client_secret=os.environ["GOOGLE_OAUTH_CLIENT_SECRET"],
        token_uri="https://oauth2.googleapis.com/token",
    )
    creds.refresh(Request())
    return build("drive", "v3", credentials=creds)


def _find_parent_folder_id(svc) -> str:
    """ค้นหา parent folder ด้วยชื่อ แทนที่จะใช้ ID ตรงๆ"""
    result = svc.files().list(
        q="name='แผนพินัยกรรม' and mimeType='application/vnd.google-apps.folder'",
        fields="files(id, name)"
    ).execute()
    files = result.get("files", [])
    if files:
        print(f"[DRIVE] Found parent: {files[0]['name']} ({files[0]['id']})")
        return files[0]["id"]
    # fallback ใช้ DRIVE_FOLDER_ID
    return DRIVE_FOLDER_ID


def create_folder(nickname: str, date_str: str) -> str:
    svc = _drive()
    parent_id = _find_parent_folder_id(svc)
    folder_name = f"ประเมินความเสี่ยง_คุณ{nickname}_{date_str}"
    meta = {
        "name": folder_name,
        "mimeType": "application/vnd.google-apps.folder",
        "parents": [parent_id],
    }
    f = svc.files().create(body=meta, fields="id").execute()
    print(f"[DRIVE] folder created: {folder_name}")
    return f["id"]


def upload_docx(path: str, filename: str, folder_id: str) -> str:
    svc = _drive()
    meta = {"name": filename, "parents": [folder_id]}
    media = MediaFileUpload(
        path,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    f = svc.files().create(
        body=meta, media_body=media, fields="id",
        supportsAllDrives=True
    ).execute()
    print(f"[DRIVE] uploaded: {filename}")
    return f["id"]


# ─── Claude: Score ────────────────────────────────────────────────────────────
SCORE_PROMPT = """คุณเป็นผู้ประเมินความพร้อมการคุ้มครองครอบครัว ตอบ JSON เท่านั้น ไม่มี markdown

## เกณฑ์คะแนน (ห้ามใช้ดุลยพินิจนอกเกณฑ์)

### หมวด 1 — สภาพคล่องฉุกเฉิน (น้ำหนัก 20%)
core: เงินพอรับมือ 90 วันแรก + ปรับตัว 1 ปี + หนี้ที่ยังเดินต่อ
- มีเงินสด ≥ 12 เดือน + มีประกันคุ้มครองหนี้ → 9-10
- มีเงินสด 6-11 เดือน → 6-8
- มีเงินสด 3-5 เดือน → 4-5
- มีเงินสด < 3 เดือน → 2-3
- ไม่รู้/ไม่ระบุ → 2
หมายเหตุ: หนี้บ้านมักมีประกันคุ้มครองหนี้แนบมา ไม่ต้องนับซ้ำ

### หมวด 2 — ความคุ้มครองชีวิต (น้ำหนัก 25%)
core: HLV = รายได้ต่อเดือน × 12 × ปีที่เหลือถึงอายุ 60
เปรียบเทียบทุนประกันรวมกับ HLV (ไม่ใช่หนี้):
- ทุนประกัน ≥ 70% HLV → 9-10
- ทุนประกัน ≥ 40% HLV → 6-8
- ทุนประกัน ≥ 10% HLV หรือ มีประกันแต่ไม่รู้ทุน → 4-5
- ไม่มีประกัน + breadwinner คนเดียว → 1-2
- ไม่มีประกันเลย → 1

### หมวด 3 — การจัดการมรดก (น้ำหนัก 25%)
core: ทรัพย์สินไปถึงคนที่ต้องการ ไม่มีปัญหา ไม่ทะเลาะ
- มีพินัยกรรม + POA + Living Will → 9-10
- มีพินัยกรรม + อย่างน้อย 1 ใน POA/Living Will → 6-8
- มีแค่พินัยกรรม → 4-5
- ไม่มีอะไรเลย → 1-2

### หมวด 4 — การดูแลลูก (น้ำหนัก 20%)
core: คนดูแลถูกคน คุณภาพชีวิตเหมือนเดิม เงินถูกใช้ตรงจุดประสงค์
- ตกลงไว้แล้ว + แยก Person/Money Guardian ชัดเจน → 9-10
- ตกลงไว้แล้ว แต่ไม่แยก guardian → 5-7
- มีคนในใจ แต่ยังไม่ตกลง → 3-4
- ไม่มีแผนเลย → 1-2
- ไม่มีลูก → null (ไม่นำไป weight)

### หมวด 5 — ความพร้อมเอกสาร (น้ำหนัก 10%)
core: คนข้างหลังมีคู่มือปฏิบัติทันที รู้ว่าต้องทำอะไร
- ครอบครัวรู้ที่เก็บเอกสาร + เคยคุยแผนฉุกเฉินแล้ว → 9-10
- ครอบครัวรู้ที่เก็บเอกสาร แต่ยังไม่เคยคุย → 5-6
- ตัวเองรู้ แต่ครอบครัวไม่รู้ + ไม่เคยคุย → 3-4
- ไม่มีใครรู้เลย → 1-2

## กฎเหล็ก
- ห้ามให้ > 5 ถ้าไม่มีเอกสาร/แผนในหมวดนั้น
- "ไม่ระบุ" ≠ 1 เสมอ — ใช้บริบทประกอบ
- ห้ามให้ < 3 ถ้ามีบางอย่างอยู่แล้ว แม้จะไม่ครบ
- ห้ามปลอบใจด้วยคะแนน — ถ้าข้อมูลบอกว่าอ่อนแอ คะแนนต้องสะท้อนความจริง

## Overall Score (weighted)
คำนวณจาก weight ตามหมวด:
- ถ้าไม่มีลูก (score_4 = null) → กระจาย weight ของหมวด 4 ไปให้หมวด 1,2,3,5 ตามสัดส่วน
- Overall 8.0-10 → ความเสี่ยงต่ำ
- Overall 6.0-7.9 → ความเสี่ยงปานกลาง
- Overall 4.0-5.9 → ความเสี่ยงสูง
- Overall 0-3.9 → ความเสี่ยงสูงมาก

ตอบ JSON:
{"score_1":int,"score_2":int,"score_3":int,"score_4":int|null,"score_5":int,
"overall":float,"risk_level":"ต่ำ"|"ปานกลาง"|"สูง"|"สูงมาก",
"hlv_estimate":int,
"gaps":["gap1","gap2",...]}"""


REPORT_PROMPT = """คุณเป็นนักเขียนรายงานการเงินสไตล์ storyselling

## สไตล์รวม
- เพื่อนที่เป็นผู้เชี่ยวชาญ ไม่สั่งสอน ไม่ขู่ ไม่แนะนำ
- ใช้ตัวเลขจริงจาก chatlog เท่านั้น ไม่แต่ง
- ห้ามใช้ "ควร" "น่าจะ" "ลองพิจารณา" "เริ่มต้น" หรือชี้นำการกระทำใดๆ
- Metaphor/อุปมาอุปมัย ใช้ได้ แต่เฉพาะเมื่อทำให้จับต้องได้ขึ้นจริง ไม่ใช้พร่ำเพื่อ

---

## โครงสร้าง (เขียนครบทุกท่อน ไม่ข้าม)

### เปิดเรื่อง (3-4 ประโยค)
- เลือกคำพูดจริงของลูกค้า 1 ประโยคเดียวที่ impact ที่สุด จากช่วงถามเรื่อง พินัยกรรม / เอกสารอยู่ที่ไหน / คุยแผนฉุกเฉินไหม
- เลือกแง่มุมเดียว ไม่รวมหลายคำพูด
- ต่อด้วยการฉายภาพสถานการณ์จากคำพูดนั้น สื่อว่า "ยังไม่ได้เตรียม" โดยไม่พูดตรงๆ
- Style: เรียบ ตรง มีน้ำหนัก ให้ลูกค้ารู้สึกว่าเขียนให้ตัวเอง

### หมวด 1 — สภาพคล่องฉุกเฉิน (4-5 ประโยค)
core: "90 วันแรก ครอบครัวจะอยู่ได้ไหม?"
- ห้าม quote คำพูดลูกค้าในหมวด — ใช้ narrative ล้วน
- ค่าใช้จ่ายที่รอไม่ได้ทันทีหลังเกิดเหตุ
- ค่าใช้จ่ายปรับตัวอย่างน้อย 1 ปี
- หนี้สินที่ยังเดินต่อแม้เจ้าของจะไม่อยู่แล้ว
- ใช้ตัวเลขจริง เงินสด runway หนี้ ฉายภาพให้เห็น
- Metaphor ใช้ได้ถ้าช่วยให้ตัวเลข abstract จับต้องได้ขึ้น
จบด้วย: [คะแนน: X/10]

### หมวด 2 — ความคุ้มครองชีวิต (4-5 ประโยค)
core: "มูลค่าเศรษฐกิจของคุณ 20 ปีข้างหน้า มีอะไร replace ไหม?"
- ห้าม quote คำพูดลูกค้าในหมวด — ใช้ narrative ล้วน
- HLV = รายได้ทั้งชีวิตที่ครอบครัวจะสูญเสีย ไม่ใช่แค่ปีนี้
- ฉายภาพช่องว่างระหว่างทุนประกันที่มี (หรือไม่มี) กับ HLV
- Metaphor ใช้ได้ถ้าช่วยให้ขนาดของ HLV รู้สึกได้จริง
จบด้วย: [คะแนน: X/10]

### หมวด 3 — การจัดการมรดก (4-5 ประโยค)
core: "ทรัพย์สินจะไปถึงคนที่คุณต้องการจริงๆ ไหม?"
- ห้าม quote คำพูดลูกค้าในหมวด — ใช้ narrative ล้วน
- ถ้าไม่มีพินัยกรรม กฎหมายแบ่งแทน ไม่ใช่คุณ
- กระบวนการศาลใช้เวลา ครอบครัวเข้าถึงทรัพย์สินได้ยาก
- ความขัดแย้งที่อาจตามมาในครอบครัว
- เรื่องราวพูดแทนตัวเองได้ ไม่จำเป็นต้องใช้ Metaphor
จบด้วย: [คะแนน: X/10]

### หมวด 4 — การดูแลลูก (4-5 ประโยค)
core: "ลูกจะได้คุณภาพชีวิตเหมือนเดิมไหม?"
- ห้าม quote คำพูดลูกค้าในหมวด — ใช้ narrative ล้วน
- ถ้าไม่มีแผน ศาลตัดสินว่าใครดูแลลูก ไม่ใช่คุณ
- คนดูแลใช่คนที่เลือกไหม เลี้ยงดูแบบที่อยากให้ไหม
- เงินที่เตรียมไว้จะถูกใช้ตรงจุดประสงค์ไหม
- Style: narrative อารมณ์ เบามือ ไม่กล่าวหา
- ข้ามหมวดนี้ทั้งหมดถ้าไม่มีลูก
จบด้วย: [คะแนน: X/10]

### หมวด 5 — ความพร้อมเอกสาร (4-5 ประโยค)
core: "คนข้างหลังมีคู่มือปฏิบัติทันทีไหม?"
- ห้าม quote คำพูดลูกค้าในหมวด — ใช้ narrative ล้วน
- ฉายภาพวันที่เกิดเหตุ คนข้างหลังต้องทำอะไรท่ามกลางความเจ็บปวด
- รู้ว่าเอกสารอยู่ที่ไหน เข้าถึงได้ทันทีไหม
- รู้ว่าต้องโทรหาใคร แจ้งอะไรก่อน หรือต้องเดาเอง
จบด้วย: [คะแนน: X/10]

### Overall
[Overall: X.X/10 — ระดับความเสี่ยง]
[พบช่องโหว่ X จุด]

### สรุป (4-5 ประโยค)
- ภาพรวมสถานการณ์ทั้งหมด ไม่แนะนำ ไม่ชี้นำ
- ย้อนให้เห็นภาพรวมทั้ง 5 หมวด ว่าตอนนี้สถานการณ์เป็นยังไง
- ปิดด้วยประโยคที่ทำให้ลูกค้า "รู้สึกบางอย่าง" โดยไม่ต้องบอก
- Style: ภาษาตรง เรียบ มีน้ำหนัก ไม่ใช้ Metaphor จังหวะนี้ต้องการความชัด ไม่ใช่ความสวย

---

## กฎเหล็ก
- ห้ามแนะนำวิธีแก้ปัญหาทุกกรณีในทุกท่อน
- ความยาวรวม 500-600 คำ (ไม่นับ tag)
- plain text ไม่มี markdown ไม่มี header — ยกเว้น tag [คะแนน:] [Overall:] [พบช่องโหว่]
- ตอบเฉพาะเนื้อหา ไม่มีคำอธิบาย
- ต้องเขียนครบทุกหมวดจนจบ ห้ามหยุดกลางหมวด
- เรียกลูกค้าว่า "คุณ[ชื่อ]" เสมอ ห้ามเรียกชื่อเปล่าๆ"""


async def _claude(system: str, user: str, max_tokens: int = 800) -> str:
    headers = {
        "x-api-key": ANTHROPIC_API_KEY,
        "anthropic-version": "2023-06-01",
        "content-type": "application/json",
    }
    payload = {
        "model": CLAUDE_MODEL,
        "max_tokens": max_tokens,
        "system": system,
        "messages": [{"role": "user", "content": user}],
    }
    async with httpx.AsyncClient(timeout=60) as client:
        resp = await client.post(ANTHROPIC_API_URL, headers=headers, json=payload)
        resp.raise_for_status()
        data = resp.json()

    # Debug: log stop reason
    stop_reason = data.get("stop_reason", "unknown")
    usage = data.get("usage", {})
    print(f"[CLAUDE] stop_reason={stop_reason} tokens={usage}")

    content = data.get("content", [])
    if not content:
        raise ValueError(f"Claude returned no content (stop_reason={stop_reason})")

    text = content[0].get("text", "").strip()
    if not text:
        raise ValueError(f"Claude returned empty text (stop_reason={stop_reason})")

    return text


async def gen_scores(chatlog: dict) -> dict:
    msgs = "\n".join(f"{m['role'].upper()}: {m['content']}" for m in chatlog.get("messages", []))
    raw = await _claude(SCORE_PROMPT, f"Chatlog:\n{msgs}")
    raw = raw.replace("```json", "").replace("```", "").strip()
    return json.loads(raw)


async def gen_content(chatlog: dict, scores: dict) -> str:
    msgs = "\n".join(f"{m['role'].upper()}: {m['content']}" for m in chatlog.get("messages", []))
    has_children = scores.get("score_4") is not None

    prompt = (
        f"Chatlog:\n{msgs}\n\n"
        f"คะแนนแต่ละหมวด (ใส่ในตำแหน่ง [คะแนน: X/10] ให้ตรงตามนี้):\n"
        f"- หมวด 1 สภาพคล่อง: {scores.get('score_1')}/10\n"
        f"- หมวด 2 ความคุ้มครองชีวิต: {scores.get('score_2')}/10 (HLV ประมาณ {scores.get('hlv_estimate', 0):,} บาท)\n"
        f"- หมวด 3 การจัดการมรดก: {scores.get('score_3')}/10\n"
        f"- หมวด 4 การดูแลลูก: {scores.get('score_4') if has_children else 'ไม่มีลูก — ข้ามหมวดนี้'}\n"
        f"- หมวด 5 ความพร้อมเอกสาร: {scores.get('score_5')}/10\n"
        f"- Overall: {scores.get('overall')}/10 — {scores.get('risk_level')}\n"
        f"- พบช่องโหว่: {len(scores.get('gaps', []))} จุด\n\n"
        f"เขียน narrative ตามโครงสร้างใน prompt\n"        f"สำคัญ: ต้องเขียนครบทุกหมวดจนจบ ห้ามหยุดกลางหมวด ห้ามตัดทิ้ง"
    )
    return await _claude(REPORT_PROMPT, prompt, max_tokens=2500)


# ─── Build .docx ──────────────────────────────────────────────────────────────
def _color(hex_str: str) -> RGBColor:
    r, g, b = int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16)
    return RGBColor(r, g, b)


ACCENT  = "1A3A5C"
DARK    = "1A1A2E"
MUTED   = "666688"
FAINT   = "AAAAAA"
RED     = "C0392B"
ORANGE  = "D35400"
YELLOW  = "B7950B"
GREEN   = "1E8449"


def risk_color(overall: float) -> str:
    if overall >= 8: return GREEN
    if overall >= 6: return YELLOW
    if overall >= 4: return ORANGE
    return RED


def risk_label(overall: float) -> str:
    if overall >= 8: return "ความเสี่ยงต่ำ"
    if overall >= 6: return "ความเสี่ยงปานกลาง"
    if overall >= 4: return "ความเสี่ยงสูง"
    return "ความเสี่ยงสูงมาก"


def score_color(s: int) -> str:
    if s >= 8: return GREEN
    if s >= 6: return YELLOW
    if s >= 4: return ORANGE
    return RED


def _parse_content_sections(content: str):
    """แยก narrative และ tag คะแนนออกจากกัน"""
    import re
    sections = []
    # แยกแต่ละ paragraph
    parts = re.split(r'(\[คะแนน:[^\]]+\]|\[Overall:[^\]]+\]|\[พบช่องโหว่[^\]]+\])', content)
    for part in parts:
        part = part.strip()
        if not part:
            continue
        if re.match(r'\[คะแนน:', part):
            sections.append(("score", part))
        elif re.match(r'\[Overall:', part):
            sections.append(("overall", part))
        elif re.match(r'\[พบช่องโหว่', part):
            sections.append(("gaps", part))
        else:
            sections.append(("text", part))
    return sections


def build_docx(nickname: str, date_th: str, content: str, scores: dict, gaps: list) -> str:
    doc = Document()

    # Page margins (A4)
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width  = Cm(21.0)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(3.0)

    def para(text, bold=False, size=11, color=DARK, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=6, italic=False):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size  = Pt(size)
        run.font.color.rgb = _color(color)
        return p

    def label(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(text.upper())
        run.font.size  = Pt(8)
        run.font.color.rgb = _color(FAINT)
        run.font.bold  = True
        # เส้นบนบาง
        pf = p.paragraph_format
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        top = OxmlElement("w:top")
        top.set(qn("w:val"), "single")
        top.set(qn("w:sz"), "4")
        top.set(qn("w:space"), "1")
        top.set(qn("w:color"), "E8E6DF")
        pBdr.append(top)
        pPr.append(pBdr)

    overall = scores.get("overall", 0)

    # ─── Header ───────────────────────────────────────────────────────────────
    para("PAYAT FINANCIAL PLANNING", bold=True, size=9, color=ACCENT, space_after=2)
    para("รายงานประเมินความพร้อมการคุ้มครองครอบครัว", bold=True, size=16, color=DARK, space_after=2)
    para(f"สำหรับคุณ{nickname}  ·  {date_th}", size=10, color=FAINT, space_after=14)

    # ─── Narrative ────────────────────────────────────────────────────────────
    label("บทวิเคราะห์")
    for line in content.split("\n"):
        if line.strip():
            para(line.strip(), size=11, color=MUTED, space_after=6)

    # ─── Score ────────────────────────────────────────────────────────────────
    label("คะแนนรวม")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f"{overall:.1f}")
    r1.bold = True
    r1.font.size = Pt(36)
    r1.font.color.rgb = _color(risk_color(overall))
    r2 = p.add_run(" / 10")
    r2.font.size = Pt(14)
    r2.font.color.rgb = _color(FAINT)

    para(risk_label(overall), bold=True, size=11, color=risk_color(overall), space_after=10)

    # ─── Score breakdown ──────────────────────────────────────────────────────
    label("คะแนนรายหมวด")
    rows = [
        ("สภาพคล่องฉุกเฉิน", scores.get("score_1")),
        ("ความคุ้มครองชีวิต", scores.get("score_2")),
        ("การจัดการมรดก",     scores.get("score_3")),
        ("การดูแลลูก",        scores.get("score_4")),
        ("ความพร้อมเอกสาร",  scores.get("score_5")),
    ]
    for lbl, s in rows:
        if s is None:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r1 = p.add_run(f"{lbl:<18}")
        r1.font.size = Pt(10)
        r1.font.color.rgb = _color(MUTED)
        r2 = p.add_run(f"  {s}/10")
        r2.bold = True
        r2.font.size = Pt(10)
        r2.font.color.rgb = _color(score_color(s))

    # ─── Gaps ─────────────────────────────────────────────────────────────────
    label(f"พบช่องโหว่ {len(gaps)} จุดที่ควรรับทราบ")
    for g in gaps:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Pt(12)
        r = p.add_run(f"—  {g}")
        r.font.size = Pt(10)
        r.font.color.rgb = _color(MUTED)

    # ─── CTA ──────────────────────────────────────────────────────────────────
    label("ขั้นตอนถัดไป")
    para(
        "รายงานนี้เป็นภาพรวมเบื้องต้น ช่องโหว่ที่พบต้องการแผนที่ออกแบบเฉพาะสำหรับครอบครัวของคุณ "
        "คุณพยัตพร้อมนั่งคุยโดยตรงเพื่อวางแผนที่ครอบคลุมทุกด้าน",
        size=10, color=MUTED, space_after=8
    )
    para("คุณพยัต จิรสุวรรณพงศ์", bold=True, size=11, color=ACCENT, space_after=2)
    para("นักวางแผนการเงิน · ที่ปรึกษากฎหมาย", size=9, color=FAINT, space_after=2)
    para("นัดคุยฟรี 30 นาที  ·  LINE: @payat  ·  payat.jira@gmail.com", size=9, color=ACCENT, space_after=2)
    para("บริการวางแผนคุ้มครองครอบครัวเต็มรูปแบบ เริ่มต้น 1,990 บาท", size=9, color=FAINT, space_after=10)

    # Footer note
    para(
        "รายงานนี้จัดทำโดย AI ภายใต้การดูแลของคุณพยัต · ข้อมูลทั้งหมดเป็นความลับ · ไม่ใช่คำแนะนำทางกฎหมายหรือการเงิน",
        size=8, color=FAINT, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=20
    )

    # Save
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    return tmp.name


# ─── Notify คุณพยัต ──────────────────────────────────────────────────────────
def notify_payat(text: str):
    import urllib.request
    data = json.dumps({
        "to": PAYAT_USER_ID,
        "messages": [{"type": "text", "text": text}]
    }).encode()
    req = urllib.request.Request(
        "https://api.line.me/v2/bot/message/push",
        data=data,
        headers={
            "Authorization": f"Bearer {LINE_CHANNEL_ACCESS_TOKEN}",
            "Content-Type": "application/json",
        },
        method="POST",
    )
    urllib.request.urlopen(req)
    print(f"[NOTIFY] LINE sent to Payat")


# ─── Main Pipeline ────────────────────────────────────────────────────────────
async def run(chatlog: dict):
    nickname = chatlog.get("nickname") or "ลูกค้า"
    date_str = datetime.now(TZ_THAI).strftime("%d%m%Y")
    date_th  = datetime.now(TZ_THAI).strftime("%-d/%m/%Y")
    filename = f"{nickname}_{date_str}.docx"

    print(f"[AGENT] Starting: {nickname}")

    scores  = await gen_scores(chatlog)
    print(f"[AGENT] Overall={scores.get('overall')} Risk={scores.get('risk_level')}")

    content = await gen_content(chatlog, scores)
    gaps    = scores.get("gaps", [])

    docx_path = build_docx(nickname, date_th, content, scores, gaps)

    try:
        folder_id = create_folder(nickname, date_str)
        file_id   = upload_docx(docx_path, filename, folder_id)
        print(f"[AGENT] Done → {filename} ({file_id})")

        # แจ้งคุณพยัต
        notify_payat(
            f"📋 รายงานใหม่พร้อมแล้วครับ\n\n"
            f"👤 ลูกค้า: คุณ{nickname}\n"
            f"📊 Overall: {scores.get('overall', 0):.1f}/10 ({scores.get('risk_level', '')})\n"
                f"⚠️ ช่องโหว่: {len(gaps)} จุด\n\n"
            f"📁 ดูใน Drive: ประเมินความเสี่ยง_คุณ{nickname}_{date_str}\n"
            f"📧 Email ลูกค้า: {chatlog.get('email', '-')}"
        )
        return file_id
    finally:
        import os as _os
        if _os.path.exists(docx_path):
            _os.unlink(docx_path)
